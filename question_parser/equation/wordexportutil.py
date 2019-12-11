import html
import re
import subprocess
from io import BytesIO
from xml.etree.ElementTree import Element

import requests
from django.conf import settings
from django.core.files.storage import default_storage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # PyCharm报错，运行没问题，不用修复
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Cm, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml.etree import _Element
from lxml.html import HtmlElement
from pyquery import PyQuery
from requests import RequestException

from question_parser.equation.aft import AftQuestion
from question_parser.equation.wordconverter import WordConverter, EquationConvertError
from question_parser.utils.misc import decode_base64

converter = WordConverter()

IMG_SIZE_ZOOM_FACTOR = 0.58  # 图片尺寸转为word尺寸的缩放引子


def str2bool(s: str):
    return s.lower() == 'true'


def brace_digit(content):
    return re.sub(r'([\^_])(\d)', r'\1{\2}', content)


def over_to_frac(content):
    """
    格式转换
    {a\\overb} -> \\frac{a}{b}
    1.  分界符{} 分组
    2.  记录左边界索引，遇到右边界弹出左边界索引，
        遇到\\over 开始查找右边界，碰到右边界弹出左边界索引, 最后一个左边界索引小于\\over索引替换数据

    :param content:
    :return:
    """
    s = ''
    left = []
    content_list = list(filter(None, re.split(r'([{}])', content)))
    content_length = len(''.join(content_list)) + 1
    over_idx = content_length
    for i in content_list:
        s += i
        if '{' == i:
            left.append(len(s) - 1)

        elif '}' == i:
            idx = left and left.pop() or 0
            if idx < over_idx:
                temp = s[idx:len(s) + 1]
                s = s.replace(temp, re.sub(r'{(.*?)\\over(?![a-zA-Z])(.*)}$', r'{\\frac{\1}{\2}}', temp))
                over_idx = content_length
        elif '\\over' in i:
            over_idx = len(s)

    return s or content


def atop_to_double_backlash(content):
    """
    \atop 转化为 反斜线
    :param content:
    :return:
    """
    return content.replace('\\atop', '\\\\')


class WordExportUtil(object):
    """
    HTML导出Word工具类。

    已知问题：
        1、不支持背景和前景颜色
        2、表格自动转换为自适应表格，丢失表格宽度等设定
        3、不支持合并单元格
    """

    def __init__(self):
        self._tmp_files = []

    def get_cm(self, px: int):
        """
        获取厘米
        :param px:
        :return:
        """
        return Cm(px / 28.346)

    def get_pt(self, px: float):
        """
        px->pt
        :param px:
        :return:
        """
        return Pt(int(px * 0.75))

    def _append_equation(self, p: Paragraph, latex: str):
        """
        增加Inline公式
        :param p:
        :param latex:
        :return:
        """

    def _get_style(self, document: Document, style_name: str):
        """
        只有保存在文档(template.docx)中的样式才能够被使用。
        确切的说，是在word的样式窗格中的"文档中的样式"的列表中的内容才可以使用，也就是只有document.styles中的才能被使用。

        将word默认样式添加到文档中的样式的方法：
            1、键入文字
            2、在样式窗格中选择样式
            3、删除步骤1中键入的文字
            4、确保 "文档中的样式"包含了刚刚应用的样式
            5、保存

        特别说明：绝对不可应用非微软word自带或windows自带字体之外的任何字体。
        :param document:
        :param style_name:
        :return:
        """
        if style_name in document.styles:
            return document.styles[style_name]
        else:
            return None

    def mini_trim(self, content):
        """
        微部调整数据
        :param content:
        :return:
        """
        content = brace_digit(content)
        content = over_to_frac(content)
        content = atop_to_double_backlash(content)
        return content

    @property
    def content_side_width(self):
        """
        内容栏宽度
        :return:
        """
        from extension.exports.exporter import WordExporter
        return (
                WordExporter.page_width -
                WordExporter.page_margin[1] -
                WordExporter.page_margin[3] -
                WordExporter.qrcode_side_width
        )

    def _render_span(self, p: Paragraph, pq: PyQuery,
                     bold=False, italic=False, strike=False, underline=False, font_size=None,
                     sub=False, sup=False):
        """
        转换span
        change 19.5.3
            公式转换错误，则直接用图片
        :param pq:
        :return:
        """
        try:
            if pq.attr('data-latex'):  # 公式
                omml_str = converter.to_omml(self.mini_trim(pq.attr('data-latex')))
                omml_str = omml_str.replace('<m:oMath',
                                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')

                pq(p._element).append(omml_str)
                return
            if pq.has_class("math-tex"):  # 公式
                if pq.attr('data-latex'):
                    omml_str = pq.attr('data-latex')
                else:
                    omml_str = html.unescape(pq.html()) if pq.html() is not None else ''
                omml_str = omml_str.replace(r'\(', '').replace(r'\)', '')
                omml_str = converter.to_omml(self.mini_trim(omml_str))

                omml_str = omml_str.replace('<m:oMath',
                                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')

                pq(p._element).append(omml_str)
                return

            # 阿凡题公式
            if pq.has_class('afanti-latex'):
                metadata = AftQuestion(pq).parse_element()
                if metadata.startswith('^') or metadata.startswith('_'):
                    last_ele = pq(p._element).children()[-1]
                    metadata = last_ele.text[-1] + metadata
                    last_ele.text = last_ele.text[:-1]

                omml_str = converter.to_omml(self.mini_trim(metadata))
                omml_str = omml_str.replace('<m:oMath',
                                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')

                pq(p._element).append(omml_str)
                return
        except EquationConvertError:
            img = PyQuery('img', pq)
            self._render_img(p, img)
            return

        bold = any(
            [
                bold,
                self._get_pq_style(pq, 'font-weight') == 'bold',
                self._get_pq_style(pq, 'font-weight') == 'bolder'
            ]
        )
        italic = any(
            [
                italic,
                self._get_pq_style(pq, 'font-style') == 'italic'
            ]
        )
        strike = any([strike, self._get_pq_style(pq, 'text-decoration') == 'line-through',
                      self._get_pq_style(pq, 'text-decoration-line') == 'line-through']
                     )
        underline = any(
            [
                underline,
                self._get_pq_style(pq, 'text-decoration') == 'underline',
                self._get_pq_style(pq, 'text-decoration-line') == 'underline'
            ]
        )

        if self._get_pq_style(pq, 'font-size'):
            size = self._get_pq_style(pq, 'font-size')
            if size.endswith('px'):
                size = size[:-2]
                size = int(float(size))
                font_size = self.get_pt(size)
            elif size.endswith('pt'):
                size = size[:-2]
                size = float(size)
                font_size = Pt(size)
        # self.__render_inline_element(p, pq, bold=bold, italic=italic, underline=underline, font_size=font_size,
        #                              strike=strike)

        contents = pq.contents()
        for item in contents:
            if isinstance(item, (HtmlElement, _Element)):
                self._render_element(p, item, is_root=True,
                                     bold=bold, italic=italic, strike=strike,
                                     underline=underline, font_size=font_size)
                continue
            run = p.add_run(self._clear_text(item))
            self.__force_simsun(run)
            if self._get_pq_style(pq, 'font-name'):
                run.font.name = self._get_pq_style(pq, 'font-name')
            if font_size:
                run.font.size = font_size

            run.underline = underline

            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            run.font.superscript = sup
            run.font.subscript = sub

    def _render_order_list(self, p: Paragraph, pq: PyQuery):
        """
        渲染序号列表
        :param p:
        :param pq:
        :return:
        """
        contents = pq.contents()
        for item in contents:
            sub_p = p._parent.add_paragraph(style='List Number')
            self._render_children(sub_p, pq(item).contents(), skip_br=True, is_root=True)

    def _render_unorder_list(self, p: Paragraph, pq: PyQuery):
        """
        渲染无序号列表
        :param p:
        :param pq:
        :return:
        """
        contents = pq.contents()
        for item in contents:
            sub_p = p._parent.add_paragraph(style='List Bullet')
            self._render_children(sub_p, pq(item).contents(), skip_br=True, is_root=True)

    def _render_table(self, p: Paragraph, pq: PyQuery):
        """
        渲染表格
        :param p:
        :param pq:
        :return:
        """
        rows = pq('tr')
        tds = pq(rows[0])('td')
        if tds is None:
            tds = pq(rows[0])('th')
        if tds is None:
            return

        # 计算行数
        row_count = len(rows)

        # 计算列数，如果有合并，则计算未合并列数
        col_count = 0
        for td in tds:
            pq_td = PyQuery(td)
            if pq_td.attr('colspan'):
                col_count += int(pq_td.attr('colspan'))
            else:
                col_count += 1

        if isinstance(p._parent, _Cell):
            table = p._parent.add_table(row_count, col_count)
        else:
            table = p.part.document.add_table(row_count, col_count)

        spans = {}  # 合并行属性。键值：row:col，值: 行(纵)合并值

        # 生成
        row_index = 0
        for row in rows:
            tds = pq(row)('td')
            if tds is None:
                tds = pq(row)('th')
            col_index = 0
            for td in tds:
                pq_td = PyQuery(td)
                # 过滤掉被合并的单元格
                col_index = self._skip_to_next_col(col_index, pq_td, row_index, spans)
                try:
                    cell = table.cell(row_index, col_index)
                    cell_p = cell.paragraphs[0]

                    self._render_children(cell_p, PyQuery(td).contents(), is_root=True)
                except IndexError:
                    pass

                # 处理合并行.
                # 将被合并的单元格，以 [行:列]=[行被合并:列被合并] 形式记录到spans中。
                if pq_td.attr('rowspan'):
                    rowspan = int(pq_td.attr('rowspan'))
                    # spans['{}:{}'.format(row_index, col_index)] = True
                    cell = table.cell(row_index, col_index)
                    # 同时合并了列
                    if pq_td.attr("colspan"):
                        # 标记被合并单元格为被合并
                        colspan = int(pq_td.attr('colspan'))
                        for i in range(0, rowspan):
                            for j in range(0, colspan):
                                key = '{}:{}'.format((i + row_index), col_index + j)
                                spans[key] = '{}:{}'.format(i == 0, True)  # 当前行设置为行合并，其他设置非行合并
                    else:
                        colspan = 1
                        for i in range(0, rowspan):
                            key = '{}:{}'.format(row_index + i, col_index)
                            spans[key] = '{}:{}'.format(True, False)  # 当前行设置为行合并，其他设置非行合并
                    # 合并单元格
                    c = table.cell(row_index + rowspan - 1, col_index + colspan - 1)
                    cell.merge(c)
                    col_index += colspan - 1
                elif pq_td.attr('colspan'):
                    cell = table.cell(row_index, col_index)
                    colspan = int(pq_td.attr('colspan'))
                    for i in range(0, colspan):
                        key = '{}:{}'.format(row_index, col_index + i)
                        spans[key] = '{}:{}'.format(False, True)

                    cell.merge(table.cell(row_index, col_index + colspan - 1))
                    col_index += colspan - 1

                col_index += 1
            row_index += 1
        table.style = self._get_style(p.part.document, "Table")

    def _skip_to_next_col(self, col_index, pq_td, row_index, spans):
        while True:
            span_sets = spans.get('{}:{}'.format(row_index, col_index), None)
            if not span_sets:  # 无合并
                break
            span_sets = span_sets.split(':')
            # 被合并。列数增加，以定位到准确位置
            if str2bool(span_sets[0]) or str2bool(span_sets[1]):
                col_index += 1
            else:
                break
        return col_index

    def adjust_pic_width(self, pic, col1_width):
        """
        调整图片宽度
        :param pic:
        :param col1_width:
        :return:
        """
        if pic.width > col1_width:
            pic.height = int(pic.height / (pic.width / col1_width))
            pic.width = int(col1_width)

    def _convert_svg_to_jpg(self, target_file_name):
        if not target_file_name.endswith(".svg"):
            return target_file_name
        jpg_target_file_name = target_file_name[:-4] + ".jpg"
        cmd = f'{settings.IMAGE_MAGICK_BIN} {target_file_name}  {jpg_target_file_name}'
        subprocess.call(cmd, shell=True)
        return jpg_target_file_name

    def _render_img(self, p: Paragraph, pq: PyQuery):
        """
        渲染图片
        :param p:
        :param pq:
        :return:
        """
        from django.conf import settings
        src = pq.attr('src')
        if src is None:
            return
        width = self._get_pq_style(pq, 'width')
        col1_width = Cm(self.content_side_width)
        if width:
            digit_array = re.findall(r'\d+(?:\.\d+)*', width)
            if len(digit_array):
                width = float(digit_array[0])
                width = min(
                    self.get_cm(int(width * IMG_SIZE_ZOOM_FACTOR)),
                    col1_width
                )

        if src.startswith("http"):

            src = src[len(settings.MEDIA_URL):]
        elif src.startswith('/media/'):
            src = src[len('/media/'):]
        if src.startswith('/'):
            src = src[1:]

        target_file_name = default_storage.path(src)

        if not default_storage.exists(target_file_name):
            # target_file_name = default_storage.path('tmp/export/word/' + src[src.rindex('/') + 1:])
            if src.startswith('data:image'):
                idx = src.index(',')
                stream = BytesIO(decode_base64(src[idx + 1:].encode('ascii')))
                pic = p.add_run().add_picture(stream, width)
                self.adjust_pic_width(pic, col1_width)
            else:
                try:
                    resp = requests.get(settings.MEDIA_URL + src, stream=True, timeout=1)
                    if resp.status_code == 200:
                        default_storage.save(target_file_name, resp.raw)
                        target_file_name = self._convert_svg_to_jpg(target_file_name)
                        pic = p.add_run().add_picture(target_file_name, width)  # 设置图片大小
                        self.adjust_pic_width(pic, col1_width)
                    else:
                        p.add_run("MISS IMG")
                        print(f"缺少图片:{src}")
                except RequestException:
                    pass
        else:
            try:
                target_file_name = self._convert_svg_to_jpg(target_file_name)
                pic = p.add_run().add_picture(target_file_name, width)  # 设置图片大小
                self.adjust_pic_width(pic, col1_width)
            except UnrecognizedImageError:
                print(f"缺少图片:{src}")
                p.add_run("MISS IMG")

    def _get_pq_style(self, pq, style: str):
        """
        获取指定PyQuery的值。优先直接设置的属性值，然后是style中的值

        百分比值将被忽略
        :param pq:
        :param style:
        :return:
        """
        if pq.attr(style):
            return pq.attr(style)
        styles = pq.attr('style')
        if styles is None:
            return None
        styles = styles.split(';')
        for item in styles:
            item = item.strip()
            if item.startswith(style + ":"):
                value = item[len(style) + 1:].strip()
                return None if value.endswith("%") else value
        return None

    def __force_simsun(self, run: Run):
        """
        避免带圈数字被识别为Times New Roman，导致重叠。强制使用宋体。
        :param run:
        :return:
        """
        symbols = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨']
        for symbol in symbols:
            if symbol in run.text:
                run.font.name = '宋体'
                return

    def __render_inline_element(self, p: Paragraph, pq: PyQuery, bold=False, italic=False, sub=False, sup=False,
                                underline=False, font_size=None, strike=False):
        """
        渲染行内元素
        :param p: 段落
        :param pq: 带渲染元素
        :param bold: 加粗
        :param italic: 斜体
        :param sub: 下标
        :param sup: 上标
        :param underline: 下划线
        :param font_size:默认字号 9pt，小五号
        :return:
        """
        for item in pq.contents():
            text = item.text if isinstance(item, (HtmlElement, _Element)) else item
            if isinstance(item, (HtmlElement, _Element)):
                self._render_element(
                    p, item,
                    bold=bold, italic=italic,
                    underline=underline, strike=strike,
                    sup=sup,
                    sub=sub,
                    font_size=font_size
                )
                continue
            run = p.add_run(text)
            self.__force_simsun(run)
            run.underline = underline
            run.bold = bold
            run.italic = italic
            run.font.superscript = sup
            run.font.subscript = sub
            if font_size:
                run.font.size = font_size
            run.font.strike = strike

    def _render_element(self, p: Paragraph, element: str or Element, is_root=False,
                        bold=False, italic=False, strike=False, underline=False, font_size=None, sup=False, sub=False):
        """
        转换html节点到word
        :param element:
        :return:
        """
        if isinstance(element, str):
            run = p.add_run(self._clear_text(element))
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
            run.font.underline = underline
            run.font.subscript = sub
            run.font.superscript = sup
            if font_size:
                run.font.size = font_size
            self.__force_simsun(run)
            return
        pq = PyQuery(element)
        if pq.is_('p'):  # 不支持嵌套p，自动扁平化
            contents = pq.contents()
            align = self._get_pq_style(pq, 'text-align')

            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if is_root:
                self._render_children(p, contents)
            else:
                sub_p = p._parent.add_paragraph()

                if align == 'center':
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                self._render_children(sub_p, contents)
        elif pq.is_('u'):  # 下划线
            self.__render_inline_element(p, pq, underline=True, bold=bold, italic=italic, strike=strike,
                                         font_size=font_size, sub=sub, sup=sup)
        elif pq.is_('strong') or pq.is_('b'):  # 加粗
            self.__render_inline_element(p, pq, underline=underline, bold=True, italic=italic, strike=strike,
                                         font_size=font_size, sub=sub, sup=sup)
        elif pq.is_('i') or pq.is_('em'):  # 斜体
            self.__render_inline_element(p, pq, underline=underline, bold=bold, italic=True, strike=strike,
                                         font_size=font_size, sub=sub, sup=sup)
        elif pq.is_('sub'):  # 下标
            self.__render_inline_element(p, pq, underline=underline, bold=bold, italic=italic, strike=strike,
                                         font_size=font_size, sub=True, sup=sup)
        elif pq.is_('sup'):  # 上标
            self.__render_inline_element(p, pq, underline=underline, bold=bold, italic=italic, strike=strike,
                                         font_size=font_size, sub=sub, sup=True)
        elif pq.is_('var'):  # 老公式
            self.__render_inline_element(p, pq, underline=underline, bold=bold, italic=True, strike=strike,
                                         font_size=font_size, sub=sub, sup=sup)
        elif pq.is_('span'):
            self._render_span(p, pq, bold=bold, italic=italic, strike=strike, underline=underline, font_size=font_size)
        elif pq.is_("br"):
            p.add_run().add_break()
        elif pq.is_("div"):
            # sub_p = p._parent.add_paragraph()
            p.add_run().add_break()
            self._render_children(p, pq.contents())
        elif pq.is_('ul'):
            self._render_unorder_list(p, pq)
        elif pq.is_('ol'):
            self._render_order_list(p, pq)
        elif pq.is_('table'):
            self._render_table(p, pq)
        elif pq.is_('img'):  # 图片
            self._render_img(p, pq)
        elif element.tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            sub_p = p._parent.add_paragraph()
            self.__render_inline_element(
                sub_p, pq, bold=True, font_size=Pt(12),
                underline=underline, italic=True, strike=strike,
                sub=sub, sup=sup
            )
        else:
            sub_p = p._parent.add_paragraph()
            contents = pq.contents()
            self._render_children(sub_p, contents)

    def _clear_text(self, text: str) -> str:
        return text.replace("\n", "")  # .strip()

    def _clear_tmp_files(self):
        """
        清理临时文件
        :return:
        """
        for f in self._tmp_files:
            default_storage.delete(f)
        self._tmp_files.clear()

    def _render_children(self, p: Paragraph, children, skip_br=False, is_root=False):
        """
        渲染子节点
        :param p:
        :param children:
        :return:
        """
        for item in children:
            if skip_br and isinstance(item, (HtmlElement, _Element)) and item.tag == 'br':
                continue
            self._render_element(p, item, is_root=is_root)

    @staticmethod
    def clear_first_p(content):
        """
        清除第一个p标签，避免不必要的换行
        :param content:
        :return:
        """
        content = content.strip()
        if content.startswith("<div>"):
            try:
                next_p = content.index("<div>", 5)
            except ValueError:
                next_p = 0
            try:
                end_p = content.index('</div>', 5)
            except ValueError:
                end_p = 0
            if end_p > next_p:  # </p> 在 <p> 后标明有最外层是<p>，而不是一组并列的<p>，可以安全移除
                content = content[5:-6].strip()
            else:
                content = content[5:end_p] + content[end_p + 6:]
        if content.startswith("<p>"):
            try:
                next_p = content.index("<p>", 3)
            except ValueError:
                next_p = 0
            try:
                end_p = content.index('</p>', 3)
            except ValueError:
                end_p = 0
            if end_p > next_p:  # </p> 在 <p> 后标明有最外层是<p>，而不是一组并列的<p>，可以安全移除
                content = content[3:-4]
            else:  # 并列一组p，把第一个移除
                content = content[3:end_p] + content[end_p + 4:]

        return content

    def render_contents(self, p: Paragraph or _Cell, body: str, flat_p=False):
        """
        渲染内容
        :param flat_p: 扁平化p标签。但是如果内容是一组并排的p，则此参数不起作用。如<p>..</p><p>..</p><p>..</p><p>..</p>，则此参数被设置为False
        :param p:
        :param body:
        :return:
        """
        if not body:
            return
        pq = PyQuery(body.replace('\r', '').replace('\n', '').replace('\t', ''))
        contents = pq.contents()
        # is_first = True
        is_partial_p = True  # 并排了一组p
        for item in contents:
            if not isinstance(item, _Element) or item.tag != 'p':
                is_partial_p = False
                break
        if is_partial_p:
            flat_p = False
        for item in contents:
            self._render_element(p, item, is_root=flat_p)
            # if is_first:
            #     pq_i = PyQuery(item)
            #     # 避免第一个是不必要的p
            #     if pq_i.is_("p"):
            #         child_contents = pq_i.contents()
            #         for child in child_contents:
            #             self._render_element(p, child)
            #     else:
            #         self._render_element(p, item)
            # else:
            #     self._render_element(p, item)
            # is_first = False

        # self._clear_tmp_files()


word_exporter = WordExportUtil()
