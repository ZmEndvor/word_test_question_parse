import html
import os
import subprocess
from logging import getLogger
from typing import Callable
from xml.etree import ElementTree

import shortuuid
from docx import Document as new_document
from docx.document import Document
from docx.oxml import nsmap, CT_Tbl, CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from question_parser.equation.dwml import omml
from question_parser.equation.mathtype import CT_R2, CT_OMath, CT_OMathPara
from question_parser.equation.mathtype.mathtype import MathTypeParser

logger = getLogger("venus")


class EquationConvertError(Exception):
    """
    公式转换错误
    """
    pass


class WordConverter(object):
    """
    转换Word中的公式.

    转换latex到omml依赖texmath
    """

    def __init__(self):
        self._file_name = ''
        self._parser = None
        self._handler = None

    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def to_latex(self, name_or_object: str or object, handler: Callable[[bytes], str] = None) -> list:
        """
        转换word文档。每个块元素(段落，表格)作为作为列表相中的一列


        :param: handler 文件保存回调函数。将word中图片数据，保存到文件系统中，并返回相对于MEDIA_ROOT的相对路径
        :return:
        :rtype: list
        """
        self._handler = handler
        if isinstance(name_or_object, str) and not os.path.exists(name_or_object):
            raise FileNotFoundError()

        self._parser = MathTypeParser()

        document = new_document(name_or_object)
        result = []
        for item in self.iter_block_items(document):
            if isinstance(item, Paragraph):
                result.append(self._convert_paragraph(item))
            elif isinstance(item, Table):
                result.append(self._convert_table(item))
        return result

    def _convert_paragraph(self, p: Paragraph):
        children = list(p._element)
        run_lst = []
        for child in children:

            if isinstance(child, CT_R2):
                run = Run(child, p)
                if 'ProgID="Equation.DSMT' in run.element.xml:
                    v = self._convert_equation(run)
                    if len(v) == 1:
                        run_lst.append(v)
                    else:
                        run_lst.append(r'<span class="math-tex" data-latex="{0}">\({0}\)</span>'.format(html.escape(v)))
                else:
                    run_lst.append(self._convert_run(run))
            elif isinstance(child, (CT_OMath, CT_OMathPara)):
                run_lst.append(self._convert_omath(child))
            elif isinstance(child, CT_Tbl):
                run_lst.append(self._convert_table(child))
        run_lst.append(self._convert_omath(p._element))
        return ''.join(run_lst)

    def _convert_equation(self, run):
        """
        转换MathType公式
        :param run:
        :return:
        """
        oleobject = run.element.embed_object.find(r'{%s}%s' % (nsmap['o'], 'OLEObject'))
        rid = oleobject.attrib[r'{%s}%s' % (nsmap['r'], 'id')]

        buf = run.part.rels[rid].target_part.blob

        self._parser._debug = True
        try:
            r = self._parser.parse(buf)
        except Exception as e:

            raise e
        return r

    def _convert_omath(self, element: ElementTree.Element or str):
        """
        转换OMath 公式

        依赖dwml转换。如有问题，可以转而使用temath
        :return:
        """

        if isinstance(element, str):
            om = omml.load_string(element)
        elif isinstance(element, CT_OMathPara):
            xml = ElementTree.tostring(element, encoding='utf8', method='xml')
            om = omml.load_string(xml)
        elif isinstance(element, CT_OMath):
            root = ElementTree.fromstring(
                '''<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:ns1="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></m:oMathPara>''')
            root.append(ElementTree.fromstring(element.xml))
            xml = ElementTree.tostring(root, encoding='utf8', method='xml')
            om = omml.load_string(xml)
        else:
            return ''
        result = []
        for o in om:
            result.append(r'<span class="math-tex" data-latex="{0}">\({0}\)</span>'.format(html.escape(o.latex)))
        return ''.join(result)

    def _convert_table(self, table: Table):
        """
        转换表格
        :param table:
        :return:
        """
        row_count = len(table.rows)
        col_count = len(table.columns)
        result = ['<table>']
        for row in range(row_count):
            result.append('<tr>')
            for col in range(col_count):
                result.append("<td>")
                cell = table.cell(row, col)
                # result.append(''.join(self.convert_(cell.paragraphs)))
                child_result = []
                for item in self.iter_block_items(cell):
                    if isinstance(item, Paragraph):
                        child_result.append(self._convert_paragraph(item))
                    elif isinstance(item, Table):
                        child_result.append(self._convert_table(item))
                result.append('<br />'.join(child_result))
                result.append("</td>")
            result.append("</tr>")
        result.append("</table>")
        return ''.join(result)

    def _convert_run(self, run):
        """
        转换Run
        """
        result = []
        formatter = '{}'
        if run.font.bold:
            formatter = '<b>' + formatter + "</b>"
        if run.font.italic:
            formatter = '<i>' + formatter + "</i>"
        if run.font.underline:
            formatter = '<u>' + formatter + "</u>"
        if run.font.subscript:
            formatter = '<sub>' + formatter + "</sub>"
        if run.font.superscript:
            formatter = '<sup>' + formatter + "</sup>"
        if run.font.math:
            print("found OMath.", end='')
        result.append(formatter.format(run.text))
        # 内嵌图片处理
        for shape in run.element.drawing_lst:
            inline = list(shape)[0]
            if not hasattr(inline, 'graphic'):
                continue

            graphic_data = inline.graphic.graphicData
            rid = graphic_data.pic.blipFill.blip.embed
            rel_obj = run.part.rels[rid]
            target_ref = rel_obj.target_ref  # media/filename
            target_obj = rel_obj.target_part
            # 文件内容 target_.obj.blob
            if self._handler:
                target_ref = self._handler(target_obj.blob, target_obj.content_type)
            result.append(target_ref)

        if run.element.pict is not None:
            imagedata = run.element.find(r'.//{%s}%s' % (nsmap['v'], 'imagedata'))
            if imagedata is not None:
                rid = imagedata.attrib[r'{%s}%s' % (nsmap['r'], 'id')]
                rel_obj = run.part.rels[rid]
                target_ref = rel_obj.target_ref  # media/filename
                target_obj = rel_obj.target_part
                # 文件内容 target_.obj.blob
                if self._handler:
                    target_ref = self._handler(target_obj.blob, target_obj.content_type)
                result.append(target_ref)
        result.append(self._convert_omath(run.element))
        return ''.join(result)

    # ################ 转换为 Word ###########
    def to_omml(self, latex: str) -> str:
        """
        转换latex到omml

        依赖texmat转换
        :param latex:
        :return:
        """
        input_file = '/tmp/{}'.format(shortuuid.uuid())

        with open(input_file, 'w') as f:
            f.write(latex)
        try:
            out = subprocess.check_output(' '.join(['texmath', '-f', 'tex', '-t', 'omml', '--inline', input_file]),
                                          shell=True, timeout=1)
            return out.decode('utf-8')
        except subprocess.CalledProcessError as e:
            logger.error(str(e) + "  " + latex)
            raise EquationConvertError(latex)
        except Exception as e:
            logger.error(e)
            raise EquationConvertError(latex)
        finally:
            os.remove(input_file)


if __name__ == '__main__':
    converter = WordConverter()
    r = converter.to_latex(r"/Users/YuanXu/Downloads/tbl.docx")  # 2017年11月16日高中数学作
    print('\r\n'.join(r))
    converter.to_omml(r"\frac{-b\pm\sqrt{b^{2}-4ac}}{2a}")
    print(r)
    # doc = new_document()
    # p = doc.add_paragraph()
    # pic = p.add_run().add_picture("/Users/yuanxu/Downloads/个人资料.png")
    # print(pic.width)
    # pic.width = int(pic.width / 2)
    # print(pic.width)
